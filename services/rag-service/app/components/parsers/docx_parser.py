"""DOCX document parser using python-docx."""

from typing import Any, Dict, List

from app.components.parsers.base import BaseParser, ParsedDocument
from app.components.parsers.factory import ParserFactory


class DocxParser(BaseParser):
    """Parser for Microsoft Word documents (.docx)."""

    name: str = "docx"
    description: str = "Parser for Microsoft Word documents (.docx) with paragraph and table extraction"
    supported_extensions: List[str] = ["docx"]

    @classmethod
    def get_config_schema(cls) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "extract_tables": {
                    "type": "boolean",
                    "default": True,
                    "description": "Extract tables as text",
                },
                "extract_headers": {
                    "type": "boolean",
                    "default": True,
                    "description": "Extract document headers/footers",
                },
                "paragraph_separator": {
                    "type": "string",
                    "default": "\n\n",
                    "description": "Separator between paragraphs",
                },
                "preserve_styles": {
                    "type": "boolean",
                    "default": False,
                    "description": "Include style information in metadata",
                },
            },
        }

    async def parse(
        self,
        content: bytes,
        filename: str,
        extract_tables: bool = True,
        extract_headers: bool = True,
        paragraph_separator: str = "\n\n",
        preserve_styles: bool = False,
        **kwargs,
    ) -> ParsedDocument:
        """Parse DOCX content.

        Args:
            content: Raw DOCX content as bytes
            filename: Original filename
            extract_tables: Whether to extract tables
            extract_headers: Whether to extract headers/footers
            paragraph_separator: Separator between paragraphs
            preserve_styles: Include style info in metadata

        Returns:
            ParsedDocument with extracted content
        """
        try:
            from docx import Document
            from io import BytesIO
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. "
                              "Install with: pip install python-docx")

        # Open document from bytes
        doc = Document(BytesIO(content))

        # Extract metadata
        metadata = {
            "filename": filename,
            "format": "docx",
        }

        # Get core properties
        core_props = doc.core_properties
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.subject:
            metadata["subject"] = core_props.subject
        if core_props.created:
            metadata["created"] = str(core_props.created)
        if core_props.modified:
            metadata["modified"] = str(core_props.modified)

        # Extract paragraphs
        paragraphs = []
        headings = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

                # Track headings
                if para.style and para.style.name.startswith("Heading"):
                    level = self._get_heading_level(para.style.name)
                    headings.append({
                        "level": level,
                        "text": text,
                    })

        if headings:
            metadata["headings"] = headings

        # Extract tables if requested
        tables = []
        table_texts = []

        if extract_tables and doc.tables:
            for i, table in enumerate(doc.tables):
                table_data = self._extract_table(table, i)
                tables.append(table_data)

                # Convert table to text
                table_text = self._table_to_text(table_data)
                if table_text:
                    table_texts.append(table_text)

        # Combine all text
        all_text_parts = paragraphs + table_texts
        full_text = paragraph_separator.join(all_text_parts)

        # Count stats
        metadata["paragraph_count"] = len(paragraphs)
        metadata["table_count"] = len(tables)

        return ParsedDocument(
            content=full_text,
            metadata=metadata,
            tables=tables if tables else None,
        )

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        if style_name == "Heading":
            return 1

        # Try to extract number from "Heading 1", "Heading 2", etc.
        parts = style_name.split()
        if len(parts) == 2 and parts[1].isdigit():
            return int(parts[1])

        return 1

    def _extract_table(self, table, index: int) -> Dict[str, Any]:
        """Extract table data."""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            rows.append(cells)

        return {
            "index": index,
            "rows": len(rows),
            "cols": len(rows[0]) if rows else 0,
            "content": rows,
        }

    def _table_to_text(self, table_data: Dict[str, Any]) -> str:
        """Convert table data to readable text."""
        if not table_data.get("content"):
            return ""

        rows = table_data["content"]
        text_rows = []

        for row in rows:
            text_rows.append(" | ".join(row))

        return "\n".join(text_rows)


# Register with factory
ParserFactory.register("docx", DocxParser)
